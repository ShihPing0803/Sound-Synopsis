from flask import Flask, render_template, request, jsonify, send_file
from scipy.io.wavfile import write
import openai
from docx import Document
from langchain_community.llms import OpenAI
from langchain_openai import OpenAI
from langchain import OpenAI
import os
import firebase_admin
from firebase_admin import credentials, auth
import azure.cognitiveservices.speech as speechsdk
from io import BytesIO
import io
import math
from openai import OpenAI
from opencc import OpenCC
from werkzeug.utils import secure_filename
from langchain_core.documents import Document
from PyPDF2 import PdfReader
from docx import Document
from doc2docx import convert

openai.api_key = os.getenv('OPENAI_API_KEY')
app = Flask(__name__)

client = OpenAI()
# Global variables
fs = 44100  # Sample rate
filename = "output.wav"
is_recording = False

class NamedBytesIO(io.BytesIO):
    name = 'transcript.wav'

def get_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def get_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def generate_title(text):
    response = client.chat.completions.create(
        model ="gpt-4o",
        messages=[
            {"role": "system", "content": "依據原文給定標題，以文字的語言為主，標題在十個字以內精簡明瞭。結果請以原文語言顯示"},
            {"role": "user", "content": text}
        ],
        max_tokens=60
    )
    title = response.choices[0].message.content
    return title

def speech_to_text():
    audio_file_path="./output.wav"
    audio_file= open(audio_file_path, "rb")
    transcript = openai.Audio.transcribe("whisper-1", audio_file)
    text = transcript.to_dict()['text']

    return text

def text_summmarization(modify_text):
    response = client.chat.completions.create(
        model ="gpt-4o",
        messages=[
            {"role": "system", "content": "依據原文進行摘要，結果請以原文語言顯示"},
            {"role": "user", "content": modify_text}
        ]
    )
    print(response)
    summary = response.choices[0].message.content
    return summary

def text_modify(original_text):
    response = client.chat.completions.create(
        model ="gpt-4o",
        messages=[
            {"role": "system", "content": "依據原文加上適當的標點符號，並且去掉冗言贅字，其餘請勿更動原內容，結果請以原文語言顯示"},
            {"role": "user", "content": original_text}
        ]
    )
    print(response)
    summary = response.choices[0].message.content
    return summary

def bullet_point(modify_text):
    response = client.chat.completions.create(
        model ="gpt-4o",
        messages=[
            {"role": "system", "content": "依據原文列點說明，結果請以原文語言顯示"},
            {"role": "user", "content": modify_text}
        ]
    )
    bullet_point = response.choices[0].message.content
    print(bullet_point)
    return bullet_point
                   
def writeDoc(text, modify, result, bullet_point_text):
    filename = "./transcript.docx"
    doc = Document()
    doc.add_heading("原文", level=3)
    doc.add_paragraph(text)
    doc.add_heading("修飾文字", level=3)
    doc.add_paragraph(modify)
    doc.add_heading("總結", level=3)
    doc.add_paragraph(result)
    doc.add_heading("列點說明", level=3)
    doc.add_paragraph(bullet_point_text)
    doc.save(filename)
    return filename

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/download_transcript', methods = ['GET'])
def download_transcript():
    filepath = './transcript.docx'
    # 检查文件是否存在
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    else:
        return jsonify({'error': '文件未找到'}), 404
    
@app.route('/process_audio', methods=['POST'])
def process_audio():
    audio_file = request.files['audio']
    if audio_file:
        file_size = len(audio_file.read())
        audio_file.seek(0) 

        max_size = 20 * 1024 * 1024
        if file_size > max_size:
            num_parts = math.ceil(file_size / max_size)
            texts = []

            for i in range(num_parts):
                part_stream = BytesIO(audio_file.read(max_size))
                part_stream.name = f'transcript_part{i}.wav'
                
                transcript_part = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=part_stream,
                    response_format='text'
                )
                texts.append(transcript_part)
            transcript = " ".join(texts)
        else:
            audio_stream = BytesIO(audio_file.read())
            audio_stream.name = 'transcript.wav'
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_stream,
                response_format='text'
            )

        cc = OpenCC('s2t')
        text = cc.convert(transcript)
        title = generate_title(text)
        modify = text_modify(text)
        result = text_summmarization(modify)
        bullet_point_text = bullet_point(modify)
        writeDoc(text, modify, result, bullet_point_text)
        title = cc.convert(title)
        modify = cc.convert(modify)
        result = cc.convert(result)
        bullet_point_text = cc.convert(bullet_point_text)

        return jsonify({'title': title,'original':text,'modify' : modify,'summary': result,'bullet_point':bullet_point_text})
    return jsonify({'error'}), 400

@app.route('/process_text',methods = ['POST'])
def process_text():
    transcript = ""
    file = request.files['file']
    filename = secure_filename(file.filename)
    file_path = os.path.join('./', filename)
    file.save(file_path)
    if filename.endswith('.pdf'):
        transcript += get_text_from_pdf(file_path)
        print(transcript)
    elif filename.endswith(".txt"):
         with open(file_path, 'r', encoding='utf-8') as file:
            transcript += file.read()
            print(transcript)
    elif filename.endswith('.docx'):
        transcript += get_text_from_docx(file_path)
        print(transcript)
    elif filename.endswith('.doc'):
        output_path = os.path.splitext(file_path)[0] + "_output.docx"
        if not os.path.exists(output_path):
            convert(file_path,output_path)
            os.remove(file_path)
        transcript += get_text_from_docx(output_path)
        print(transcript)

    print(transcript)
    cc = OpenCC('s2t')
    text = cc.convert(transcript)
    title = generate_title(text)
    modify = text_modify(text)
    result = text_summmarization(modify)
    bullet_point_text = bullet_point(modify)
    writeDoc(text, modify, result, bullet_point_text)
    title = cc.convert(title)
    modify = cc.convert(modify)
    result = cc.convert(result)
    bullet_point_text = cc.convert(bullet_point_text)

    return jsonify({'title': title,'original':text,'modify' : modify,'summary': result,'bullet_point':bullet_point_text})

if __name__ == '__main__':
    app.run(debug=True, host='localhost')
